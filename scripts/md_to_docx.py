"""Convert a Markdown file to a polished .docx.

Handles the subset of markdown used in this project:
- H1/H2/H3 headings
- Paragraphs
- Bulleted (-) and numbered (1.) lists
- Tables (GitHub pipe syntax)
- Fenced code blocks (```)
- Inline **bold**, *italic*, `code`, and [link](url)
- Horizontal rules (---)
- Blockquotes (>)
- Task-list items ([ ] and [x])

Usage:  python scripts/md_to_docx.py <input.md> <output.docx>
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# --- Style palette (amber accent + Cognizant blue) ---
COLOR_H1 = RGBColor(0x1E, 0x40, 0xAF)   # deep blue
COLOR_H2 = RGBColor(0x2E, 0x63, 0xC9)
COLOR_H3 = RGBColor(0x4B, 0x5C, 0x8B)
COLOR_ACCENT = RGBColor(0xB4, 0x54, 0x09)   # amber-brown
COLOR_MUTED = RGBColor(0x5C, 0x6A, 0x8C)
COLOR_LINK = RGBColor(0x1E, 0x40, 0xAF)
COLOR_CODE = RGBColor(0x2A, 0x2A, 0x2A)
COLOR_CODE_BG = "F1F3F7"
COLOR_TABLE_HEADER_BG = "1E40AF"
COLOR_TABLE_ROW_ALT = "F5F7FB"


# --- inline parsing helpers ---

_INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*)"       # **bold**
    r"|(?<!\*)\*([^*]+)\*"   # *italic*
    r"|(`[^`]+`)"            # `code`
    r"|(\[[^\]]+\]\([^)]+\))"  # [link](url)
)


def add_runs(paragraph, text: str) -> None:
    """Walk a markdown-ish string, adding runs with the right formatting."""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            r = paragraph.add_run(chunk[1:-1]); r.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = COLOR_ACCENT
            _shade_run(r, "EEF2F7")
        elif chunk.startswith("["):
            # [text](url) — render text with link colour
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk).groups()
            _add_hyperlink(paragraph, url, label)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _shade_run(run, hex_fill: str) -> None:
    """Give an inline run a background shade (Word's 'w:shd' element)."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    rpr.append(shd)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1E40AF"); r_pr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve"); new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _add_hr(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "94A3B8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# --- document setup ---

def build_document() -> Document:
    doc = Document()
    # Base body font
    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)

    # Widen margins slightly
    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    return doc


def _style_heading(paragraph, level: int, text: str) -> None:
    paragraph.text = ""
    r = paragraph.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(20); r.font.color.rgb = COLOR_H1
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(15); r.font.color.rgb = COLOR_H2
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(4)
    else:
        r.font.size = Pt(12.5); r.font.color.rgb = COLOR_H3
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(3)


# --- table rendering ---

def add_markdown_table(doc: Document, rows: list[str]) -> None:
    """rows is a list of raw markdown table lines including header + separator."""
    parsed = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in rows if line.strip() and not _is_table_separator(line)
    ]
    if not parsed:
        return
    n_cols = max(len(r) for r in parsed)
    parsed = [r + [""] * (n_cols - len(r)) for r in parsed]

    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            add_runs(para, cell_text)
            if r_idx == 0:
                # Header formatting
                _shade_cell(cell, COLOR_TABLE_HEADER_BG)
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10.5)
            else:
                if r_idx % 2 == 0:
                    _shade_cell(cell, COLOR_TABLE_ROW_ALT)
                for run in para.runs:
                    run.font.size = Pt(10.5)


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    return bool(stripped) and set(stripped.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")) == set()


# --- code blocks ---

def add_code_block(doc: Document, lines: list[str], language: str = "") -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_CODE
        _shade_run(r, COLOR_CODE_BG)
        # Shade the paragraph background too via border+shd
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), COLOR_CODE_BG)
        p_pr.append(shd)


# --- main parser ---

def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = build_document()

    # ---- Cover block ----
    title = doc.add_paragraph()
    r = title.add_run("Autonomous Data Engineering Factory")
    r.bold = True; r.font.size = Pt(24); r.font.color.rgb = COLOR_H1
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Demo & Test Guide")
    r.font.size = Pt(14); r.font.color.rgb = COLOR_MUTED; r.italic = True
    subtitle.paragraph_format.space_after = Pt(0)

    banner = doc.add_paragraph()
    r = banner.add_run("Powered by the Cognizant Agentic Engineering Excellence Platform")
    r.font.size = Pt(10.5); r.font.color.rgb = COLOR_ACCENT
    banner.paragraph_format.space_after = Pt(10)
    _add_hr(banner)

    # ---- Streaming parser ----
    i = 0
    in_code = False
    code_buffer: list[str] = []
    code_lang = ""
    table_buffer: list[str] = []
    # Skip the top H1 in the source (we've written our own cover)
    _skipped_first_h1 = False

    while i < len(lines):
        raw = lines[i]

        # code block toggling
        if raw.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = raw.strip().lstrip("`").strip()
                code_buffer = []
            else:
                add_code_block(doc, code_buffer, code_lang)
                in_code = False
                code_buffer = []
                code_lang = ""
            i += 1
            continue
        if in_code:
            code_buffer.append(raw)
            i += 1
            continue

        # flush a pending table
        if table_buffer and not raw.strip().startswith("|"):
            add_markdown_table(doc, table_buffer)
            table_buffer = []

        stripped = raw.strip()
        if not stripped:
            # blank line — end lists, add breathing room
            i += 1
            continue

        # HR
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            _add_hr(p)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1 and not _skipped_first_h1:
                _skipped_first_h1 = True
                i += 1
                continue
            para = doc.add_paragraph()
            _style_heading(para, level, text)
            i += 1
            continue

        # Table (start capture)
        if stripped.startswith("|"):
            table_buffer.append(raw)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            p_pr = para._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
            left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "8"); left.set(qn("w:color"), "B45409")
            pbdr.append(left); p_pr.append(pbdr)
            r = para.add_run(stripped.lstrip("> ").strip())
            r.italic = True; r.font.color.rgb = COLOR_MUTED
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*(\d+)\.\s+(.+)$", raw)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, m.group(2))
            i += 1
            continue

        # Bulleted list (incl. checklist)
        m = re.match(r"^\s*[-*]\s+(.+)$", raw)
        if m:
            content = m.group(1)
            # checklist rendering
            chk = re.match(r"^\[([ xX])\]\s+(.+)$", content)
            para = doc.add_paragraph(style="List Bullet")
            if chk:
                mark = "☒ " if chk.group(1).lower() == "x" else "☐ "
                r = para.add_run(mark); r.font.name = "Segoe UI Symbol"
                add_runs(para, chk.group(2))
            else:
                add_runs(para, content)
            i += 1
            continue

        # Fallback: paragraph
        para = doc.add_paragraph()
        add_runs(para, stripped)
        i += 1

    # flush trailing table
    if table_buffer:
        add_markdown_table(doc, table_buffer)

    doc.save(docx_path)
    print(f"[ok] {md_path.name} -> {docx_path}  ({docx_path.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("usage: python scripts/md_to_docx.py <input.md> <output.docx>")
        raise SystemExit(2)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
